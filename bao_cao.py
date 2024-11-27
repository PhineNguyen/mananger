import csv
from ttkbootstrap import ttk
from ttkbootstrap.constants import *
from tkinter import filedialog, messagebox, StringVar
from docx import Document


def create_bao_cao_tab(notebook, app):
    """
    Tạo tab Báo Cáo và Danh Sách.
    :param notebook: Notebook của ứng dụng.
    :param app: Cửa sổ chính.
    """
    # Tạo frame cho tab
    bao_cao_tab = ttk.Frame(notebook)
    notebook.add(bao_cao_tab, text="XUẤT")

    # Tạo label hướng dẫn
    label = ttk.Label(bao_cao_tab, text="Chọn loại báo cáo để in:", font=("Arial", 14))
    label.pack(pady=10)

    # Danh sách các loại báo cáo
    options = ["Danh sách sản phẩm", "Danh sách đơn hàng", "Danh sách khách hàng"]
    selected_option = StringVar(value=options[0])

    # Dropdown menu
    dropdown = ttk.Combobox(bao_cao_tab, textvariable=selected_option, values=options, state="readonly", width=30)
    dropdown.pack(pady=5)

    # Đọc dữ liệu từ CSV
    def read_csv(file_path):
        try:
            with open(file_path, mode="r", encoding="utf-8") as file:
                reader = csv.reader(file)
                data = list(reader)
                return data
        except Exception as e:
            messagebox.showerror("Lỗi", f"Không thể đọc file {file_path}: {e}")
            return []

    # Button để xuất báo cáo
    def export_report():
        selected = selected_option.get()
        file_path = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Word Document", "*.docx")]
        )
        if not file_path:
            return

        try:
            # Tạo tài liệu Word
            doc = Document()
            doc.add_heading(f"Báo cáo: {selected}", level=1)
            doc.add_paragraph("---------------------")

            if selected == "Danh sách sản phẩm":
                data = read_csv("products.csv")
                headers = data[0]
                rows = data[1:]

                # Thêm bảng vào Word
                table = doc.add_table(rows=1, cols=len(headers))
                table.style = "Table Grid"

                # Thêm header
                for idx, header in enumerate(headers):
                    table.cell(0, idx).text = header

                # Thêm dữ liệu
                for row in rows:
                    cells = table.add_row().cells
                    for idx, value in enumerate(row):
                        cells[idx].text = value

            elif selected == "Danh sách đơn hàng":
                data = read_csv("orders.csv")
                headers = data[0]
                rows = data[1:]

                # Thêm bảng vào Word
                table = doc.add_table(rows=1, cols=len(headers))
                table.style = "Table Grid"

                # Thêm header
                for idx, header in enumerate(headers):
                    table.cell(0, idx).text = header

                # Thêm dữ liệu
                for row in rows:
                    cells = table.add_row().cells
                    for idx, value in enumerate(row):
                        cells[idx].text = value

            elif selected == "Danh sách khách hàng":
                data = read_csv("customers.csv")
                headers = data[0]
                rows = data[1:]

                # Thêm bảng vào Word
                table = doc.add_table(rows=1, cols=len(headers))
                table.style = "Table Grid"

                # Thêm header
                for idx, header in enumerate(headers):
                    table.cell(0, idx).text = header

                # Thêm dữ liệu
                for row in rows:
                    cells = table.add_row().cells
                    for idx, value in enumerate(row):
                        cells[idx].text = value

            # Lưu tài liệu
            doc.save(file_path)
            messagebox.showinfo("Thành công", f"Đã lưu báo cáo tại: {file_path}")
        except Exception as e:
            messagebox.showerror("Lỗi", f"Không thể lưu báo cáo: {e}")

    export_button = ttk.Button(bao_cao_tab, text="In Báo Cáo", command=export_report, bootstyle=PRIMARY)
    export_button.pack(pady=10)

    # Giao diện hoàn chỉnh
    ttk.Label(bao_cao_tab, text="Hãy chọn loại báo cáo và bấm 'In Báo Cáo' để xuất.", font=("Arial", 10)).pack(pady=5)
